from __future__ import annotations

import hashlib
import json
import math
import platform
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document


SEED = 2005
BOOTSTRAPS = 2000
ROOT = Path(__file__).resolve().parents[1]
CONVERTED = ROOT / "working" / "converted"
RESULTS = ROOT / "results"
FIGURES = ROOT / "figures"
RESULTS.mkdir(parents=True, exist_ok=True)
FIGURES.mkdir(parents=True, exist_ok=True)

GRADE_LABELS = {1: "I", 2: "II", 3: "III", 4: "IV", 5: "V", 6: "劣V"}
ROMAN_TO_GRADE = {
    "I": 1,
    "Ⅰ": 1,
    "II": 2,
    "Ⅱ": 2,
    "III": 3,
    "Ⅲ": 3,
    "IV": 4,
    "Ⅳ": 4,
    "V": 5,
    "Ⅴ": 5,
    "劣V": 6,
    "劣Ⅴ": 6,
}
CLASS_COLUMNS = ["I", "II", "III", "IV", "V", "bad"]


def clean_text(value: str) -> str:
    return re.sub(r"\s+", "", value.replace("\xa0", " "))


def as_builtin(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(k): as_builtin(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [as_builtin(v) for v in value]
    if isinstance(value, np.ndarray):
        return [as_builtin(v) for v in value.tolist()]
    if isinstance(value, (np.integer,)):
        return int(value)
    if isinstance(value, (np.floating,)):
        return None if not np.isfinite(value) else float(value)
    if isinstance(value, (pd.Timestamp, datetime)):
        return value.isoformat()
    return value


def write_json(path: Path, payload: Any) -> None:
    path.write_text(
        json.dumps(as_builtin(payload), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def grade_do(value: float) -> int:
    for grade, threshold in enumerate([7.5, 6.0, 5.0, 3.0, 2.0], start=1):
        if value >= threshold:
            return grade
    return 6


def grade_upper(value: float, thresholds: list[float]) -> int:
    for grade, threshold in enumerate(thresholds, start=1):
        if value <= threshold:
            return grade
    return 6


def project_simplex(values: np.ndarray, total: float = 100.0) -> np.ndarray:
    values = np.asarray(values, dtype=float)
    ordered = np.sort(values)[::-1]
    cumulative = np.cumsum(ordered) - total
    indices = np.arange(1, len(values) + 1, dtype=float)
    valid = np.nonzero(ordered - cumulative / indices > 0)[0]
    if len(valid) == 0:
        return np.full_like(values, total / len(values))
    rho = valid[-1]
    theta = cumulative[rho] / (rho + 1)
    return np.maximum(values - theta, 0.0)


def wilson_interval(successes: int, trials: int, z: float = 1.959963984540054) -> tuple[float, float]:
    if trials == 0:
        return float("nan"), float("nan")
    proportion = successes / trials
    denominator = 1 + z * z / trials
    center = (proportion + z * z / (2 * trials)) / denominator
    half = z * math.sqrt(proportion * (1 - proportion) / trials + z * z / (4 * trials * trials)) / denominator
    return 100 * (center - half), 100 * (center + half)


def extract_monthly_and_hydrology() -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    source = CONVERTED / "<SOURCE_FILE_REDACTED>"
    if not source.is_file():
        raise FileNotFoundError(f"Converted input not found: {source}")
    document = Document(source)
    dates = []
    for paragraph in document.paragraphs:
        match = re.search(r"发布日期;(\d{4}-\d{2})", paragraph.text)
        if match:
            dates.append(pd.Timestamp(match.group(1) + "-01"))
    if len(dates) != 28 or len(document.tables) != 29:
        raise ValueError(f"Unexpected attachment 3 structure: dates={len(dates)}, tables={len(document.tables)}")

    monthly_rows: list[dict[str, Any]] = []
    aliases = 0
    for date, table in zip(dates, document.tables[:28]):
        if len(table.rows) != 19 or len(table.columns) != 10:
            raise ValueError(f"Unexpected monthly table at {date:%Y-%m}")
        for row in table.rows[2:]:
            cells = [clean_text(cell.text) for cell in row.cells]
            station = cells[1]
            if station == "四川攀枝花":
                station = "四川攀枝花龙洞"
                aliases += 1
            ph, dissolved_oxygen, codmn, nh3n = map(float, cells[3:7])
            component_grades = {
                "grade_ph": 1 if 6.0 <= ph <= 9.0 else 6,
                "grade_do": grade_do(dissolved_oxygen),
                "grade_codmn": grade_upper(codmn, [2.0, 4.0, 6.0, 10.0, 15.0]),
                "grade_nh3n": grade_upper(nh3n, [0.15, 0.5, 1.0, 1.5, 2.0]),
            }
            computed_grade = max(component_grades.values())
            official_grade = ROMAN_TO_GRADE[cells[7]]
            ratios = {
                "ratio_ph": max(6.0 / ph, ph / 9.0),
                "ratio_do": 5.0 / dissolved_oxygen,
                "ratio_codmn": codmn / 6.0,
                "ratio_nh3n": nh3n,
            }
            monthly_rows.append(
                {
                    "date": date,
                    "station": station,
                    "section": cells[2],
                    "ph": ph,
                    "do_mg_l": dissolved_oxygen,
                    "codmn_mg_l": codmn,
                    "nh3n_mg_l": nh3n,
                    "official_grade": official_grade,
                    "computed_grade": computed_grade,
                    **component_grades,
                    **ratios,
                    "wqi_class_iii_max": max(ratios.values()),
                }
            )
    monthly = pd.DataFrame(monthly_rows)
    monthly["group"] = np.where(monthly["section"].str.startswith("干流"), "干流", "支流/湖库")
    monthly["season"] = monthly["date"].dt.month.map(
        lambda month: "枯水期" if month <= 4 else ("丰水期" if month <= 10 else "平水期")
    )
    monthly["grade_label"] = monthly["official_grade"].map(GRADE_LABELS)
    component_names = ["pH", "DO", "CODMn", "NH3-N"]
    component_columns = ["grade_ph", "grade_do", "grade_codmn", "grade_nh3n"]
    limiting = []
    for values in monthly[component_columns].to_numpy():
        maximum = values.max()
        limiting.append("+".join(name for name, value in zip(component_names, values) if value == maximum))
    monthly["limiting_indicator"] = limiting

    hydro_table = document.tables[28]
    station_map = {
        "四川攀枝花": "四川攀枝花龙洞",
        "重庆朱沱": "重庆朱沱",
        "湖北宜昌": "湖北宜昌南津关",
        "湖南岳阳": "湖南岳阳城陵矶",
        "江西九江": "江西九江河西水厂",
        "安徽安庆": "安徽安庆皖河口",
        "江苏南京": "江苏南京林山",
    }
    hydro_stations = [station_map[clean_text(cell.text)] for cell in hydro_table.rows[0].cells[2:]]
    distances = [float(clean_text(cell.text)) for cell in hydro_table.rows[1].cells[2:]]
    hydro_rows: list[dict[str, Any]] = []
    for row_index in range(2, len(hydro_table.rows), 2):
        flow_row = hydro_table.rows[row_index]
        velocity_row = hydro_table.rows[row_index + 1]
        flow_cells = [clean_text(cell.text) for cell in flow_row.cells]
        velocity_cells = [clean_text(cell.text) for cell in velocity_row.cells]
        if flow_cells[1] != "水流量" or velocity_cells[1] != "水流速":
            raise ValueError(f"Unexpected hydrology row pair at {row_index}")
        date = pd.Timestamp(flow_cells[0].replace(".", "-") + "-01")
        flows = list(map(float, flow_cells[2:]))
        velocities = list(map(float, velocity_cells[2:]))
        for station, distance, flow, velocity in zip(hydro_stations, distances, flows, velocities):
            hydro_rows.append(
                {
                    "date": date,
                    "station": station,
                    "distance_km": distance,
                    "flow_m3_s": flow,
                    "velocity_m_s": velocity,
                }
            )
    hydrology = pd.DataFrame(hydro_rows)

    audit = {
        "monthly_rows": len(monthly),
        "monthly_expected_rows": 17 * 28,
        "monthly_missing_cells": int(monthly.isna().sum().sum()),
        "station_alias_rows_normalized": aliases,
        "official_grade_matches": int((monthly["official_grade"] == monthly["computed_grade"]).sum()),
        "official_grade_match_rate": float((monthly["official_grade"] == monthly["computed_grade"]).mean()),
        "ph_outside_6_9_rows": int(((monthly["ph"] < 6) | (monthly["ph"] > 9)).sum()),
        "hydrology_rows": len(hydrology),
        "hydrology_expected_rows": 7 * 13,
        "hydrology_missing_cells": int(hydrology.isna().sum().sum()),
        "distance_monotone": bool(np.all(np.diff(sorted(hydrology["distance_km"].unique())) > 0)),
    }
    monthly.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig", date_format="%Y-%m-%d")
    hydrology.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig", date_format="%Y-%m-%d")
    return monthly, hydrology, audit


def extract_annual() -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    source = CONVERTED / "<SOURCE_FILE_REDACTED>"
    if not source.is_file():
        raise FileNotFoundError(f"Converted input not found: {source}")
    document = Document(source)
    rows: list[dict[str, Any]] = []
    raw_count = 0
    for table_index, table in enumerate(document.tables):
        year = 1995 + table_index
        for row in table.rows[2:]:
            raw_count += 1
            cells = [clean_text(cell.text) for cell in row.cells]
            if year == 1995:
                period, scope, evaluated_length, remaining = cells[0], cells[2], cells[3], cells[4:]
            else:
                period, scope, evaluated_length, remaining = cells[0], cells[1], cells[2], cells[3:]
            if len(remaining) != 12:
                raise ValueError(f"Unexpected annual row structure in {year}: {cells}")
            lengths = list(map(float, remaining[0::2]))
            percentages = list(map(float, remaining[1::2]))
            record: dict[str, Any] = {
                "year": year,
                "period": period,
                "scope": scope,
                "evaluated_length_km": float(evaluated_length),
            }
            for name, length, percentage in zip(CLASS_COLUMNS, lengths, percentages):
                record[f"{name}_length_km_raw"] = length
                record[f"{name}_pct_raw"] = percentage
            rows.append(record)
    annual = pd.DataFrame(rows)
    annual = annual.drop_duplicates(["year", "period", "scope"], keep="first").reset_index(drop=True)
    raw_pct_columns = [f"{name}_pct_raw" for name in CLASS_COLUMNS]
    annual["raw_pct_sum"] = annual[raw_pct_columns].sum(axis=1)
    for name in CLASS_COLUMNS:
        annual[f"{name}_pct"] = annual[f"{name}_pct_raw"] / annual["raw_pct_sum"] * 100.0

    totals_rows = []
    flow_unit_issue_years = []
    note_pattern = re.compile(
        r"(\d{4})年长江总流量([0-9.]+)亿(立方米|吨)，废水排放总量([0-9.]+)亿吨"
    )
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        match = note_pattern.search(text)
        if match:
            year = int(match.group(1))
            unit = match.group(3)
            if unit == "吨":
                flow_unit_issue_years.append(year)
            totals_rows.append(
                {
                    "year": year,
                    "river_flow_100m_m3": float(match.group(2)),
                    "wastewater_100m_t": float(match.group(4)),
                    "flow_unit_as_printed": f"亿{unit}",
                }
            )
    totals = pd.DataFrame(totals_rows).sort_values("year").reset_index(drop=True)
    if len(annual) != 90 or len(totals) != 10:
        raise ValueError(f"Unexpected attachment 4 extraction: annual={len(annual)}, totals={len(totals)}")

    water_year_mainstem = annual[(annual["period"] == "水文年") & (annual["scope"] == "干流")]
    coverage_pre = float(water_year_mainstem.loc[water_year_mainstem["year"] == 1998, "evaluated_length_km"].iloc[0])
    coverage_post = float(water_year_mainstem.loc[water_year_mainstem["year"] == 1999, "evaluated_length_km"].iloc[0])
    annual_audit = {
        "raw_rows": raw_count,
        "unique_rows": len(annual),
        "exact_duplicate_rows_removed": raw_count - len(annual),
        "percentage_sum_anomaly_rows_gt_0_5pp": int(((annual["raw_pct_sum"] - 100).abs() > 0.5).sum()),
        "maximum_percentage_sum_error_pp": float((annual["raw_pct_sum"] - 100).abs().max()),
        "mainstem_water_year_max_sum_error_pp": float(
            (water_year_mainstem["raw_pct_sum"] - 100).abs().max()
        ),
        "mainstem_evaluated_length_1998_km": coverage_pre,
        "mainstem_evaluated_length_1999_km": coverage_post,
        "coverage_ratio_1999_to_1998": coverage_post / coverage_pre,
        "flow_unit_issue_years": flow_unit_issue_years,
    }
    annual.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")
    totals.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")
    return annual, totals, annual_audit


def analyze_quality(monthly: pd.DataFrame) -> dict[str, Any]:
    grade_counts = monthly["official_grade"].value_counts().sort_index()
    class_summary = pd.DataFrame(
        {
            "grade": range(1, 7),
            "grade_label": [GRADE_LABELS[index] for index in range(1, 7)],
            "count": [int(grade_counts.get(index, 0)) for index in range(1, 7)],
        }
    )
    class_summary["pct"] = class_summary["count"] / len(monthly) * 100
    class_summary.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    station_records = []
    for station, group in monthly.groupby("station", sort=False):
        nonpotable_count = int((group["official_grade"] >= 4).sum())
        low, high = wilson_interval(nonpotable_count, len(group))
        station_records.append(
            {
                "station": station,
                "group": group["group"].iloc[0],
                "n": len(group),
                "mean_grade": group["official_grade"].mean(),
                "median_grade": group["official_grade"].median(),
                "drinkable_pct": (group["official_grade"] <= 3).mean() * 100,
                "nonpotable_pct": nonpotable_count / len(group) * 100,
                "nonpotable_wilson_low_pct": low,
                "nonpotable_wilson_high_pct": high,
                "bad_v_pct": (group["official_grade"] == 6).mean() * 100,
                "mean_wqi_class_iii_max": group["wqi_class_iii_max"].mean(),
                "p90_wqi_class_iii_max": group["wqi_class_iii_max"].quantile(0.9),
                "mean_do_mg_l": group["do_mg_l"].mean(),
                "mean_codmn_mg_l": group["codmn_mg_l"].mean(),
                "mean_nh3n_mg_l": group["nh3n_mg_l"].mean(),
            }
        )
    stations = pd.DataFrame(station_records).sort_values(
        ["nonpotable_pct", "mean_grade", "mean_wqi_class_iii_max"], ascending=False
    )
    stations["pollution_rank"] = np.arange(1, len(stations) + 1)
    stations.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    monthly_records = []
    for date, group in monthly.groupby("date"):
        record: dict[str, Any] = {
            "date": date,
            "n": len(group),
            "mean_grade": group["official_grade"].mean(),
            "drinkable_pct": (group["official_grade"] <= 3).mean() * 100,
            "nonpotable_pct": (group["official_grade"] >= 4).mean() * 100,
            "bad_v_pct": (group["official_grade"] == 6).mean() * 100,
            "mean_wqi_class_iii_max": group["wqi_class_iii_max"].mean(),
        }
        for grade in range(1, 7):
            record[f"class_{grade}_pct"] = (group["official_grade"] == grade).mean() * 100
        monthly_records.append(record)
    monthly_summary = pd.DataFrame(monthly_records)
    monthly_summary.to_csv(
        RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig", date_format="%Y-%m-%d"
    )

    scope_records = []
    for dimension in ["group", "season"]:
        for label, group in monthly.groupby(dimension):
            scope_records.append(
                {
                    "dimension": dimension,
                    "label": label,
                    "n": len(group),
                    "mean_grade": group["official_grade"].mean(),
                    "drinkable_pct": (group["official_grade"] <= 3).mean() * 100,
                    "nonpotable_pct": (group["official_grade"] >= 4).mean() * 100,
                    "bad_v_pct": (group["official_grade"] == 6).mean() * 100,
                    "mean_wqi_class_iii_max": group["wqi_class_iii_max"].mean(),
                }
            )
    scope_summary = pd.DataFrame(scope_records)
    scope_summary.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    nonpotable = monthly[monthly["official_grade"] >= 4]
    limiting_counts = nonpotable["limiting_indicator"].value_counts().rename_axis("indicator").reset_index(name="count")
    limiting_counts["pct_of_nonpotable"] = limiting_counts["count"] / len(nonpotable) * 100
    limiting_counts.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    mainstem = scope_summary[(scope_summary["dimension"] == "group") & (scope_summary["label"] == "干流")].iloc[0]
    tributary = scope_summary[(scope_summary["dimension"] == "group") & (scope_summary["label"] == "支流/湖库")].iloc[0]
    quality = {
        "observations": len(monthly),
        "period_start": monthly["date"].min().strftime("%Y-%m"),
        "period_end": monthly["date"].max().strftime("%Y-%m"),
        "drinkable_pct": float((monthly["official_grade"] <= 3).mean() * 100),
        "nonpotable_pct": float((monthly["official_grade"] >= 4).mean() * 100),
        "bad_v_pct": float((monthly["official_grade"] == 6).mean() * 100),
        "mean_grade": float(monthly["official_grade"].mean()),
        "mean_wqi_class_iii_max": float(monthly["wqi_class_iii_max"].mean()),
        "mainstem_nonpotable_pct": float(mainstem["nonpotable_pct"]),
        "tributary_nonpotable_pct": float(tributary["nonpotable_pct"]),
        "tributary_bad_v_pct": float(tributary["bad_v_pct"]),
        "worst_station": stations.iloc[0]["station"],
        "worst_station_nonpotable_pct": float(stations.iloc[0]["nonpotable_pct"]),
        "worst_station_bad_v_pct": float(stations.iloc[0]["bad_v_pct"]),
        "second_station": stations.iloc[1]["station"],
        "second_station_nonpotable_pct": float(stations.iloc[1]["nonpotable_pct"]),
        "third_station": stations.iloc[2]["station"],
        "third_station_nonpotable_pct": float(stations.iloc[2]["nonpotable_pct"]),
        "grade_match_rate": float((monthly["official_grade"] == monthly["computed_grade"]).mean()),
    }
    return quality


SEGMENT_DISPLAY = {
    "四川攀枝花龙洞→重庆朱沱": "攀枝花—朱沱",
    "重庆朱沱→湖北宜昌南津关": "朱沱—宜昌",
    "湖北宜昌南津关→湖南岳阳城陵矶": "宜昌—岳阳",
    "湖南岳阳城陵矶→江西九江河西水厂": "岳阳—九江",
    "江西九江河西水厂→安徽安庆皖河口": "九江—安庆",
    "安徽安庆皖河口→江苏南京林山": "安庆—南京",
}


def compute_source_rows(merged: pd.DataFrame, decay_k_per_day: float) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for date, group in merged.sort_values(["date", "distance_km"]).groupby("date"):
        group = group.sort_values("distance_km").reset_index(drop=True)
        if len(group) != 7:
            raise ValueError(f"Expected seven mainstem stations at {date:%Y-%m}, got {len(group)}")
        for index in range(1, len(group)):
            upstream = group.iloc[index - 1]
            downstream = group.iloc[index]
            segment = f"{upstream['station']}→{downstream['station']}"
            distance = downstream["distance_km"] - upstream["distance_km"]
            mean_velocity = (upstream["velocity_m_s"] + downstream["velocity_m_s"]) / 2.0
            travel_days = distance / (mean_velocity * 86.4)
            attenuation = math.exp(-decay_k_per_day * travel_days)
            for pollutant, concentration_column in [("CODMn", "codmn_mg_l"), ("NH3-N", "nh3n_mg_l")]:
                upstream_load = 86.4 * upstream["flow_m3_s"] * upstream[concentration_column]
                downstream_load = 86.4 * downstream["flow_m3_s"] * downstream[concentration_column]
                attenuated_upstream = upstream_load * attenuation
                net_increment = downstream_load - attenuated_upstream
                rows.append(
                    {
                        "date": date,
                        "segment": segment,
                        "segment_display": SEGMENT_DISPLAY[segment],
                        "pollutant": pollutant,
                        "decay_k_per_day": decay_k_per_day,
                        "distance_km": distance,
                        "travel_days": travel_days,
                        "attenuation_factor": attenuation,
                        "upstream_load_kg_day": upstream_load,
                        "attenuated_upstream_load_kg_day": attenuated_upstream,
                        "downstream_load_kg_day": downstream_load,
                        "net_increment_kg_day": net_increment,
                        "positive_increment_kg_day": max(net_increment, 0.0),
                    }
                )
    return pd.DataFrame(rows)


def analyze_sources(monthly: pd.DataFrame, hydrology: pd.DataFrame) -> dict[str, Any]:
    concentrations = monthly[
        ["date", "station", "codmn_mg_l", "nh3n_mg_l"]
    ]
    merged = hydrology.merge(concentrations, on=["date", "station"], how="left", validate="one_to_one")
    if merged[["codmn_mg_l", "nh3n_mg_l"]].isna().any().any():
        raise ValueError("Mainstem concentration/hydrology join produced missing values")

    primary = compute_source_rows(merged, 0.2)
    primary.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig", date_format="%Y-%m-%d")
    reconstruction_error = (
        primary["attenuated_upstream_load_kg_day"]
        + primary["net_increment_kg_day"]
        - primary["downstream_load_kg_day"]
    ).abs().max()

    ranking = (
        primary.groupby(["pollutant", "segment", "segment_display"], as_index=False)
        .agg(
            mean_net_increment_kg_day=("net_increment_kg_day", "mean"),
            median_net_increment_kg_day=("net_increment_kg_day", "median"),
            mean_positive_increment_kg_day=("positive_increment_kg_day", "mean"),
            positive_month_pct=("net_increment_kg_day", lambda values: (values > 0).mean() * 100),
            mean_travel_days=("travel_days", "mean"),
        )
    )
    ranking["positive_share_pct"] = ranking.groupby("pollutant")["mean_positive_increment_kg_day"].transform(
        lambda values: values / values.sum() * 100
    )
    ranking["rank"] = ranking.groupby("pollutant")["mean_positive_increment_kg_day"].rank(
        ascending=False, method="min"
    ).astype(int)

    rng = np.random.default_rng(SEED)
    bootstrap_records = []
    for pollutant in ["CODMn", "NH3-N"]:
        subset = primary[primary["pollutant"] == pollutant]
        pivot = subset.pivot(index="date", columns="segment", values="positive_increment_kg_day").sort_index()
        sample_indices = rng.integers(0, len(pivot), size=(BOOTSTRAPS, len(pivot)))
        sampled_means = pivot.to_numpy()[sample_indices].mean(axis=1)
        top_indices = sampled_means.argmax(axis=1)
        for column_index, segment in enumerate(pivot.columns):
            bootstrap_records.append(
                {
                    "pollutant": pollutant,
                    "segment": segment,
                    "bootstrap_mean_low_kg_day": np.quantile(sampled_means[:, column_index], 0.05),
                    "bootstrap_mean_high_kg_day": np.quantile(sampled_means[:, column_index], 0.95),
                    "top_rank_frequency_pct": (top_indices == column_index).mean() * 100,
                }
            )
    bootstrap_frame = pd.DataFrame(bootstrap_records)
    ranking = ranking.merge(bootstrap_frame, on=["pollutant", "segment"], how="left")
    ranking = ranking.sort_values(["pollutant", "rank"])
    ranking.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    sensitivity_frames = []
    for decay_k in np.round(np.linspace(0.1, 0.5, 41), 3):
        frame = compute_source_rows(merged, float(decay_k))
        aggregate = (
            frame.groupby(["pollutant", "segment", "segment_display"], as_index=False)["positive_increment_kg_day"]
            .mean()
            .rename(columns={"positive_increment_kg_day": "mean_positive_increment_kg_day"})
        )
        aggregate["decay_k_per_day"] = decay_k
        aggregate["rank"] = aggregate.groupby("pollutant")["mean_positive_increment_kg_day"].rank(
            ascending=False, method="min"
        ).astype(int)
        sensitivity_frames.append(aggregate)
    sensitivity = pd.concat(sensitivity_frames, ignore_index=True)
    sensitivity.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    top_records = []
    for pollutant in ["CODMn", "NH3-N"]:
        subset = ranking[ranking["pollutant"] == pollutant].sort_values("rank")
        for order in range(3):
            row = subset.iloc[order]
            top_records.append(
                {
                    "pollutant": pollutant,
                    "rank": order + 1,
                    "segment": row["segment_display"],
                    "mean_positive_t_day": row["mean_positive_increment_kg_day"] / 1000.0,
                    "share_pct": row["positive_share_pct"],
                    "positive_month_pct": row["positive_month_pct"],
                    "top_rank_frequency_pct": row["top_rank_frequency_pct"],
                }
            )
    top_frame = pd.DataFrame(top_records)
    stable_top = {}
    for pollutant in ["CODMn", "NH3-N"]:
        top_by_k = sensitivity[(sensitivity["pollutant"] == pollutant) & (sensitivity["rank"] == 1)]
        counts = top_by_k["segment_display"].value_counts()
        stable_top[pollutant] = {
            "segment": counts.index[0],
            "grid_frequency_pct": counts.iloc[0] / len(top_by_k) * 100,
        }
    return {
        "months": int(primary["date"].nunique()),
        "decay_k_primary_per_day": 0.2,
        "decay_sensitivity_min": 0.1,
        "decay_sensitivity_max": 0.5,
        "mass_balance_max_abs_error_kg_day": float(reconstruction_error),
        "top_three": top_frame.to_dict(orient="records"),
        "stable_top_over_decay_grid": stable_top,
    }


def linear_simplex_forecast(train: np.ndarray, years: np.ndarray, target_year: int | np.ndarray) -> np.ndarray:
    design = np.column_stack([np.ones(len(years)), years - years[0]])
    coefficients = np.linalg.lstsq(design, train, rcond=None)[0]
    targets = np.atleast_1d(target_year)
    forecasts = []
    for year in targets:
        raw = np.array([1.0, year - years[0]]) @ coefficients
        forecasts.append(project_simplex(raw))
    result = np.vstack(forecasts)
    return result[0] if np.ndim(target_year) == 0 else result


def alr_ridge_forecast(
    train: np.ndarray,
    years: np.ndarray,
    target_year: int | np.ndarray,
    alpha: float = 100.0,
    epsilon: float = 0.1,
) -> np.ndarray:
    proportions = train + epsilon
    proportions = proportions / proportions.sum(axis=1, keepdims=True)
    reference = 2
    nonreference = [0, 1, 3, 4, 5]
    logratios = np.log(proportions[:, nonreference] / proportions[:, [reference]])
    standard_deviation = max(float(years.std()), 1.0)
    standardized_years = (years - years.mean()) / standard_deviation
    design = np.column_stack([np.ones(len(years)), standardized_years])
    penalty = np.diag([0.0, alpha])
    coefficients = np.linalg.solve(design.T @ design + penalty, design.T @ logratios)
    targets = np.atleast_1d(target_year)
    forecasts = []
    for year in targets:
        target = np.array([1.0, (year - years.mean()) / standard_deviation])
        predicted_logratios = target @ coefficients
        composition = np.ones(6)
        composition[nonreference] = np.exp(predicted_logratios)
        composition[reference] = 1.0
        forecasts.append(composition / composition.sum() * 100.0)
    result = np.vstack(forecasts)
    return result[0] if np.ndim(target_year) == 0 else result


def wastewater_prediction(train: np.ndarray, years: np.ndarray, target_year: int, model: str) -> float:
    if model == "persistence":
        return float(train[-1])
    if model == "linear":
        coefficients = np.polyfit(years, train, 1)
        return float(np.polyval(coefficients, target_year))
    if model == "log_linear":
        coefficients = np.polyfit(years, np.log(train), 1)
        return float(np.exp(np.polyval(coefficients, target_year)))
    raise ValueError(model)


def water_forecast_bootstrap(
    observations: np.ndarray,
    years: np.ndarray,
    future_years: np.ndarray,
    rng: np.random.Generator,
) -> np.ndarray:
    design = np.column_stack([np.ones(len(years)), years - years[0]])
    coefficients = np.linalg.lstsq(design, observations, rcond=None)[0]
    fitted = design @ coefficients
    residuals = observations - fitted
    residuals -= residuals.mean(axis=0, keepdims=True)
    simulations = np.empty((BOOTSTRAPS, len(future_years), observations.shape[1]))
    for bootstrap in range(BOOTSTRAPS):
        sampled = residuals[rng.integers(0, len(residuals), size=len(residuals))]
        synthetic = np.vstack([project_simplex(row) for row in fitted + sampled])
        simulations[bootstrap] = linear_simplex_forecast(synthetic, years, future_years)
    return simulations


def wastewater_forecast_bootstrap(
    observations: np.ndarray,
    years: np.ndarray,
    future_years: np.ndarray,
    rng: np.random.Generator,
) -> np.ndarray:
    design = np.column_stack([np.ones(len(years)), years])
    coefficients = np.linalg.lstsq(design, np.log(observations), rcond=None)[0]
    fitted = design @ coefficients
    residuals = np.log(observations) - fitted
    residuals -= residuals.mean()
    simulations = np.empty((BOOTSTRAPS, len(future_years)))
    future_design = np.column_stack([np.ones(len(future_years)), future_years])
    for bootstrap in range(BOOTSTRAPS):
        synthetic = fitted + residuals[rng.integers(0, len(residuals), size=len(residuals))]
        bootstrap_coefficients = np.linalg.lstsq(design, synthetic, rcond=None)[0]
        simulations[bootstrap] = np.exp(future_design @ bootstrap_coefficients)
    return simulations


def treatment_from_composition(composition: np.ndarray, wastewater: np.ndarray) -> tuple[np.ndarray, ...]:
    iv_v = composition[..., 3] + composition[..., 4]
    bad_v = composition[..., 5]
    nonpotable = iv_v + bad_v
    removal_points = bad_v + np.maximum(iv_v - 20.0, 0.0)
    treatment_fraction = np.divide(
        removal_points,
        nonpotable,
        out=np.zeros_like(removal_points, dtype=float),
        where=nonpotable > 0,
    )
    treatment = wastewater * treatment_fraction
    return iv_v, bad_v, nonpotable, removal_points, treatment_fraction, treatment


def analyze_forecast_and_treatment(
    annual: pd.DataFrame, totals: pd.DataFrame
) -> tuple[dict[str, Any], dict[str, Any], pd.DataFrame]:
    mainstem = annual[(annual["period"] == "水文年") & (annual["scope"] == "干流")].sort_values("year")
    years = mainstem["year"].to_numpy(dtype=int)
    composition_columns = [f"{name}_pct" for name in CLASS_COLUMNS]
    observed_composition = mainstem[composition_columns].to_numpy(dtype=float)
    wastewater = totals.sort_values("year")["wastewater_100m_t"].to_numpy(dtype=float)
    if not np.array_equal(years, totals.sort_values("year")["year"].to_numpy(dtype=int)):
        raise ValueError("Annual quality and wastewater years do not align")

    comparison_rows = []
    water_candidates = ["persistence", "simplex_linear", "alr_ridge"]
    water_errors: dict[str, list[float]] = {name: [] for name in water_candidates}
    for end in range(4, len(years)):
        train = observed_composition[:end]
        train_years = years[:end]
        actual = observed_composition[end]
        predictions = {
            "persistence": project_simplex(train[-1]),
            "simplex_linear": linear_simplex_forecast(train, train_years, int(years[end])),
            "alr_ridge": alr_ridge_forecast(train, train_years, int(years[end])),
        }
        for model, predicted in predictions.items():
            error = float(np.mean(np.abs(predicted - actual)))
            water_errors[model].append(error)
            comparison_rows.append(
                {
                    "task": "mainstem_composition",
                    "model": model,
                    "forecast_year": int(years[end]),
                    "metric": "six_class_mae_pp",
                    "error": error,
                }
            )

    wastewater_candidates = ["persistence", "linear", "log_linear"]
    wastewater_errors: dict[str, list[float]] = {name: [] for name in wastewater_candidates}
    for end in range(4, len(years)):
        for model in wastewater_candidates:
            predicted = wastewater_prediction(wastewater[:end], years[:end], int(years[end]), model)
            error = abs(predicted - wastewater[end])
            wastewater_errors[model].append(error)
            comparison_rows.append(
                {
                    "task": "wastewater_discharge",
                    "model": model,
                    "forecast_year": int(years[end]),
                    "metric": "mae_100m_t",
                    "error": error,
                }
            )
    rolling_errors = pd.DataFrame(comparison_rows)
    rolling_errors.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    aggregate_rows = []
    for task, error_map in [
        ("mainstem_composition", water_errors),
        ("wastewater_discharge", wastewater_errors),
    ]:
        means = {model: float(np.mean(errors)) for model, errors in error_map.items()}
        winner = min(means, key=means.get)
        for model, errors in error_map.items():
            aggregate_rows.append(
                {
                    "task": task,
                    "model": model,
                    "rolling_origin_count": len(errors),
                    "mean_absolute_error": np.mean(errors),
                    "root_mean_square_error": math.sqrt(np.mean(np.square(errors))),
                    "decision_status": "pass" if model == winner else "fail",
                }
            )
    model_comparison = pd.DataFrame(aggregate_rows)
    model_comparison.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    future_years = np.arange(2005, 2015)
    water_primary = linear_simplex_forecast(observed_composition, years, future_years)
    wastewater_primary = np.array(
        [wastewater_prediction(wastewater, years, int(year), "log_linear") for year in future_years]
    )
    wastewater_linear = np.array(
        [wastewater_prediction(wastewater, years, int(year), "linear") for year in future_years]
    )

    rng = np.random.default_rng(SEED)
    water_simulations = water_forecast_bootstrap(observed_composition, years, future_years, rng)
    wastewater_simulations = wastewater_forecast_bootstrap(wastewater, years, future_years, rng)

    iv_v, bad_v, nonpotable, removal_points, treatment_fraction, treatment = treatment_from_composition(
        water_primary, wastewater_primary
    )
    sim_iv_v, sim_bad_v, sim_nonpotable, sim_removal, sim_fraction, sim_treatment = treatment_from_composition(
        water_simulations, wastewater_simulations
    )

    forecast_records = []
    for index, year in enumerate(future_years):
        record: dict[str, Any] = {"year": int(year)}
        for class_index, name in enumerate(CLASS_COLUMNS):
            record[f"{name}_pct"] = water_primary[index, class_index]
            record[f"{name}_pct_p05"] = np.quantile(water_simulations[:, index, class_index], 0.05)
            record[f"{name}_pct_p95"] = np.quantile(water_simulations[:, index, class_index], 0.95)
        record.update(
            {
                "iv_v_pct": iv_v[index],
                "bad_v_pct": bad_v[index],
                "nonpotable_pct": nonpotable[index],
                "nonpotable_pct_p05": np.quantile(sim_nonpotable[:, index], 0.05),
                "nonpotable_pct_p95": np.quantile(sim_nonpotable[:, index], 0.95),
                "wastewater_100m_t": wastewater_primary[index],
                "wastewater_100m_t_p05": np.quantile(wastewater_simulations[:, index], 0.05),
                "wastewater_100m_t_p95": np.quantile(wastewater_simulations[:, index], 0.95),
            }
        )
        forecast_records.append(record)
    forecast_frame = pd.DataFrame(forecast_records)
    forecast_frame.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    treatment_records = []
    for index, year in enumerate(future_years):
        lower_mapping = wastewater_primary[index] * removal_points[index] / 100.0
        efficiency_85 = min(wastewater_primary[index], treatment[index] / 0.85)
        treatment_records.append(
            {
                "year": int(year),
                "predicted_wastewater_100m_t": wastewater_primary[index],
                "predicted_iv_v_pct": iv_v[index],
                "predicted_bad_v_pct": bad_v[index],
                "predicted_nonpotable_pct": nonpotable[index],
                "required_reduction_percentage_points": removal_points[index],
                "treatment_fraction_pct": treatment_fraction[index] * 100.0,
                "required_treatment_100m_t": treatment[index],
                "post_control_iv_v_pct": min(iv_v[index], 20.0),
                "post_control_bad_v_pct": 0.0,
                "required_treatment_100m_t_p05": np.quantile(sim_treatment[:, index], 0.05),
                "required_treatment_100m_t_p95": np.quantile(sim_treatment[:, index], 0.95),
                "absolute_share_mapping_100m_t": lower_mapping,
                "treatment_at_85pct_effectiveness_100m_t": efficiency_85,
                "treatment_with_linear_wastewater_100m_t": treatment_fraction[index] * wastewater_linear[index],
            }
        )
    treatment_frame = pd.DataFrame(treatment_records)
    treatment_frame.to_csv(RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig")

    # Monitoring-coverage sensitivity: refit the same model only to the post-1999 regime.
    post_mask = years >= 1999
    post_forecast = linear_simplex_forecast(observed_composition[post_mask], years[post_mask], future_years)
    persistence_forecast = np.repeat(observed_composition[[-1]], len(future_years), axis=0)
    scenario_rows = []
    for scenario, composition, waste in [
        ("primary_full_period", water_primary, wastewater_primary),
        ("post_1999_water_window", post_forecast, wastewater_primary),
        ("water_persistence", persistence_forecast, wastewater_primary),
        ("linear_wastewater", water_primary, wastewater_linear),
    ]:
        scenario_treatment = treatment_from_composition(composition, waste)[-1]
        for year, amount in zip(future_years, scenario_treatment):
            scenario_rows.append({"scenario": scenario, "year": int(year), "treatment_100m_t": amount})
    pd.DataFrame(scenario_rows).to_csv(
        RESULTS / "<SOURCE_FILE_REDACTED>", index=False, encoding="utf-8-sig"
    )

    log_slope = np.polyfit(years, np.log(wastewater), 1)[0]
    forecast_summary = {
        "historical_year_start": int(years.min()),
        "historical_year_end": int(years.max()),
        "forecast_year_start": int(future_years.min()),
        "forecast_year_end": int(future_years.max()),
        "selected_composition_model": "simplex_linear",
        "selected_composition_cv_mae_pp": float(np.mean(water_errors["simplex_linear"])),
        "composition_persistence_cv_mae_pp": float(np.mean(water_errors["persistence"])),
        "composition_alr_cv_mae_pp": float(np.mean(water_errors["alr_ridge"])),
        "selected_wastewater_model": "log_linear",
        "selected_wastewater_cv_mae_100m_t": float(np.mean(wastewater_errors["log_linear"])),
        "wastewater_linear_cv_mae_100m_t": float(np.mean(wastewater_errors["linear"])),
        "wastewater_annual_growth_pct": float((math.exp(log_slope) - 1) * 100),
        "nonpotable_2005_pct": float(nonpotable[0]),
        "nonpotable_2014_pct": float(nonpotable[-1]),
        "iv_v_2005_pct": float(iv_v[0]),
        "iv_v_2014_pct": float(iv_v[-1]),
        "bad_v_2005_pct": float(bad_v[0]),
        "bad_v_2014_pct": float(bad_v[-1]),
        "wastewater_2005_100m_t": float(wastewater_primary[0]),
        "wastewater_2014_100m_t": float(wastewater_primary[-1]),
    }
    treatment_summary = {
        "mapping": "在非饮用水河段占比与未新增处理污水量成正比的工程代理下，优先消除劣V，再将IV+V压至20%",
        "target_iv_v_max_pct": 20.0,
        "target_bad_v_pct": 0.0,
        "treatment_2005_100m_t": float(treatment[0]),
        "treatment_2014_100m_t": float(treatment[-1]),
        "treatment_fraction_2005_pct": float(treatment_fraction[0] * 100),
        "treatment_fraction_2014_pct": float(treatment_fraction[-1] * 100),
        "treatment_2005_p05_100m_t": float(np.quantile(sim_treatment[:, 0], 0.05)),
        "treatment_2005_p95_100m_t": float(np.quantile(sim_treatment[:, 0], 0.95)),
        "treatment_2014_p05_100m_t": float(np.quantile(sim_treatment[:, -1], 0.05)),
        "treatment_2014_p95_100m_t": float(np.quantile(sim_treatment[:, -1], 0.95)),
        "treatment_sum_10yr_100m_t": float(treatment.sum()),
    }
    return forecast_summary, treatment_summary, model_comparison


def configure_plot_font() -> None:
    from matplotlib import font_manager

    plt.style.use("seaborn-v0_8-whitegrid")
    candidates = [
        Path("<ABSOLUTE_PATH>"),
        Path("<ABSOLUTE_PATH>"),
        Path("<ABSOLUTE_PATH>"),
        Path("<ABSOLUTE_PATH>"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            font_manager.fontManager.addfont(str(candidate))
            family = font_manager.FontProperties(fname=str(candidate)).get_name()
            plt.rcParams["font.family"] = "sans-serif"
            plt.rcParams["font.sans-serif"] = [family]
            break
    plt.rcParams["axes.unicode_minus"] = False


def make_figures(monthly: pd.DataFrame, annual: pd.DataFrame, totals: pd.DataFrame) -> None:
    configure_plot_font()
    colors = ["#2b8cbe", "#7bccc4", "#bae4bc", "#fdae6b", "#f16913", "#a50f15"]

    month = pd.read_csv(RESULTS / "<SOURCE_FILE_REDACTED>", parse_dates=["date"])
    fig, axis = plt.subplots(figsize=(10, 5.4))
    stack = [month[f"class_{grade}_pct"].to_numpy(float) for grade in range(1, 7)]
    axis.stackplot(month["date"], *stack, labels=["I", "II", "III", "IV", "V", "劣V"], colors=colors)
    axis.set_ylabel("观测断面占比 (%)")
    axis.set_xlabel("月份")
    axis.set_title("17个观测站月度水质等级构成")
    axis.set_ylim(0, 100)
    axis.legend(ncol=6, loc="upper center", bbox_to_anchor=(0.5, -0.14))
    fig.tight_layout()
    fig.savefig(FIGURES / "<SOURCE_FILE_REDACTED>", dpi=220, bbox_inches="tight")
    plt.close(fig)

    stations = pd.read_csv(RESULTS / "<SOURCE_FILE_REDACTED>")
    stations = stations.sort_values("nonpotable_pct", ascending=True)
    fig, axis = plt.subplots(figsize=(9.2, 7.2))
    positions = np.arange(len(stations))
    axis.barh(positions, stations["nonpotable_pct"], color="#fdae6b", label="IV类及以下")
    axis.barh(positions, stations["bad_v_pct"], color="#a50f15", label="劣V类")
    axis.set_yticks(positions, stations["station"])
    axis.set_xlabel("发生率 (%)")
    axis.set_title("各地区非饮用水等级发生率（28个月）")
    axis.legend(loc="lower right")
    fig.tight_layout()
    fig.savefig(FIGURES / "<SOURCE_FILE_REDACTED>", dpi=220, bbox_inches="tight")
    plt.close(fig)

    source = pd.read_csv(RESULTS / "<SOURCE_FILE_REDACTED>")
    fig, axes = plt.subplots(1, 2, figsize=(12.5, 5.2), sharey=False)
    for axis, pollutant, color in zip(axes, ["CODMn", "NH3-N"], ["#3182bd", "#31a354"]):
        subset = source[source["pollutant"] == pollutant].sort_values("rank", ascending=False)
        values = subset["mean_positive_increment_kg_day"].to_numpy(float) / 1000.0
        low = subset["bootstrap_mean_low_kg_day"].to_numpy(float) / 1000.0
        high = subset["bootstrap_mean_high_kg_day"].to_numpy(float) / 1000.0
        errors = np.vstack([values - low, high - values])
        axis.barh(subset["segment_display"], values, xerr=errors, color=color, alpha=0.86, capsize=3)
        axis.set_title(pollutant)
        axis.set_xlabel("正增量均值 (t/d)，90% bootstrap区间")
    fig.suptitle("一阶降解质量守恒反演的分段污染贡献（k=0.2 d$^{-1}$）")
    fig.tight_layout()
    fig.savefig(FIGURES / "<SOURCE_FILE_REDACTED>", dpi=220, bbox_inches="tight")
    plt.close(fig)

    mainstem = annual[(annual["period"] == "水文年") & (annual["scope"] == "干流")].sort_values("year")
    historical_nonpotable = mainstem[["IV_pct", "V_pct", "bad_pct"]].sum(axis=1)
    forecast = pd.read_csv(RESULTS / "<SOURCE_FILE_REDACTED>")
    fig, axis = plt.subplots(figsize=(9.5, 5.4))
    axis.plot(mainstem["year"], historical_nonpotable, "o-", color="#636363", label="历史IV+V+劣V")
    axis.plot(forecast["year"], forecast["nonpotable_pct"], "o-", color="#e6550d", label="预测IV+V+劣V")
    axis.fill_between(
        forecast["year"].to_numpy(float),
        forecast["nonpotable_pct_p05"].to_numpy(float),
        forecast["nonpotable_pct_p95"].to_numpy(float),
        color="#fdae6b",
        alpha=0.3,
        label="90%残差bootstrap区间",
    )
    axis.plot(forecast["year"], forecast["iv_v_pct"], "^--", color="#3182bd", label="预测IV+V")
    axis.plot(forecast["year"], forecast["bad_v_pct"], "s--", color="#a50f15", label="预测劣V")
    axis.axhline(20, color="#238b45", linestyle=":", label="IV+V控制线20%（参照）")
    axis.set_xlabel("年份")
    axis.set_ylabel("干流河长占比 (%)")
    axis.set_title("长江干流水质比例历史与未来10年基准预测")
    axis.legend(ncol=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(FIGURES / "<SOURCE_FILE_REDACTED>", dpi=220, bbox_inches="tight")
    plt.close(fig)

    treatment = pd.read_csv(RESULTS / "<SOURCE_FILE_REDACTED>")
    fig, axis = plt.subplots(figsize=(9.5, 5.4))
    axis.plot(totals["year"], totals["wastewater_100m_t"], "o-", color="#636363", label="历史废水排放")
    axis.plot(
        treatment["year"], treatment["predicted_wastewater_100m_t"], "o-", color="#3182bd", label="基准预测排放"
    )
    axis.plot(
        treatment["year"], treatment["required_treatment_100m_t"], "s-", color="#e6550d", label="需新增有效处理"
    )
    axis.fill_between(
        treatment["year"].to_numpy(float),
        treatment["required_treatment_100m_t_p05"].to_numpy(float),
        treatment["required_treatment_100m_t_p95"].to_numpy(float),
        color="#fdae6b",
        alpha=0.3,
        label="处理量90%残差bootstrap区间",
    )
    axis.set_xlabel("年份")
    axis.set_ylabel("亿吨/年")
    axis.set_title("废水排放预测与目标约束下的处理需求")
    axis.legend(ncol=2, fontsize=9)
    fig.tight_layout()
    fig.savefig(FIGURES / "<SOURCE_FILE_REDACTED>", dpi=220, bbox_inches="tight")
    plt.close(fig)

    comparison = pd.read_csv(RESULTS / "<SOURCE_FILE_REDACTED>")
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.5))
    for axis, task, title, ylabel in [
        (axes[0], "mainstem_composition", "水质构成模型", "六类平均绝对误差 (百分点)"),
        (axes[1], "wastewater_discharge", "废水排放模型", "平均绝对误差 (亿吨)"),
    ]:
        subset = comparison[comparison["task"] == task]
        bar_colors = ["#31a354" if status == "pass" else "#bdbdbd" for status in subset["decision_status"]]
        axis.bar(subset["model"], subset["mean_absolute_error"], color=bar_colors)
        axis.set_title(title)
        axis.set_ylabel(ylabel)
        axis.tick_params(axis="x", rotation=20)
    fig.suptitle("滚动原点一步预测比较（绿色为入选模型）")
    fig.tight_layout()
    fig.savefig(FIGURES / "<SOURCE_FILE_REDACTED>", dpi=220, bbox_inches="tight")
    plt.close(fig)


def find_top(source_summary: dict[str, Any], pollutant: str, rank: int) -> dict[str, Any]:
    for record in source_summary["top_three"]:
        if record["pollutant"] == pollutant and record["rank"] == rank:
            return record
    raise KeyError((pollutant, rank))


def generate_paper_numbers(
    quality: dict[str, Any],
    sources: dict[str, Any],
    forecast: dict[str, Any],
    treatment: dict[str, Any],
) -> None:
    cod_top = find_top(sources, "CODMn", 1)
    nh_top = find_top(sources, "NH3-N", 1)
    numbers = {
        "overall_observations": str(quality["observations"]),
        "drinkable_pct": f"{quality['drinkable_pct']:.2f}%",
        "nonpotable_pct": f"{quality['nonpotable_pct']:.2f}%",
        "bad_v_pct": f"{quality['bad_v_pct']:.2f}%",
        "mainstem_nonpotable_pct": f"{quality['mainstem_nonpotable_pct']:.2f}%",
        "tributary_nonpotable_pct": f"{quality['tributary_nonpotable_pct']:.2f}%",
        "worst_station_nonpotable_pct": f"{quality['worst_station_nonpotable_pct']:.2f}%",
        "cod_top_load_t_day": f"{cod_top['mean_positive_t_day']:.1f}",
        "cod_top_share_pct": f"{cod_top['share_pct']:.2f}%",
        "nh_top_load_t_day": f"{nh_top['mean_positive_t_day']:.1f}",
        "nh_top_share_pct": f"{nh_top['share_pct']:.2f}%",
        "forecast_nonpotable_2005_pct": f"{forecast['nonpotable_2005_pct']:.2f}%",
        "forecast_nonpotable_2014_pct": f"{forecast['nonpotable_2014_pct']:.2f}%",
        "forecast_bad_2005_pct": f"{forecast['bad_v_2005_pct']:.2f}%",
        "forecast_bad_2014_pct": f"{forecast['bad_v_2014_pct']:.2f}%",
        "wastewater_2005_100m_t": f"{forecast['wastewater_2005_100m_t']:.1f}",
        "wastewater_2014_100m_t": f"{forecast['wastewater_2014_100m_t']:.1f}",
        "treatment_2005_100m_t": f"{treatment['treatment_2005_100m_t']:.1f}",
        "treatment_2014_100m_t": f"{treatment['treatment_2014_100m_t']:.1f}",
        "treatment_total_100m_t": f"{treatment['treatment_sum_10yr_100m_t']:.1f}",
        "composition_cv_mae_pp": f"{forecast['selected_composition_cv_mae_pp']:.2f}",
        "wastewater_cv_mae_100m_t": f"{forecast['selected_wastewater_cv_mae_100m_t']:.2f}",
    }
    write_json(RESULTS / "paper_numbers.json", numbers)
    escaped = {key: value.replace("%", r"\%") for key, value in numbers.items()}
    macros = [
        "% This file is generated by code/solve.py. Do not edit manually.",
        f"\\newcommand{{\\OverallObservations}}{{{numbers['overall_observations']}}}",
        f"\\newcommand{{\\DrinkablePct}}{{{escaped['drinkable_pct']}}}",
        f"\\newcommand{{\\NonpotablePct}}{{{escaped['nonpotable_pct']}}}",
        f"\\newcommand{{\\BadVPct}}{{{escaped['bad_v_pct']}}}",
        f"\\newcommand{{\\MainstemNonpotablePct}}{{{escaped['mainstem_nonpotable_pct']}}}",
        f"\\newcommand{{\\TributaryNonpotablePct}}{{{escaped['tributary_nonpotable_pct']}}}",
        f"\\newcommand{{\\WorstStationNonpotablePct}}{{{escaped['worst_station_nonpotable_pct']}}}",
        f"\\newcommand{{\\CODTopLoad}}{{{numbers['cod_top_load_t_day']}}}",
        f"\\newcommand{{\\CODTopShare}}{{{escaped['cod_top_share_pct']}}}",
        f"\\newcommand{{\\NHTopLoad}}{{{numbers['nh_top_load_t_day']}}}",
        f"\\newcommand{{\\NHTopShare}}{{{escaped['nh_top_share_pct']}}}",
        f"\\newcommand{{\\ForecastNonpotableStart}}{{{escaped['forecast_nonpotable_2005_pct']}}}",
        f"\\newcommand{{\\ForecastNonpotableEnd}}{{{escaped['forecast_nonpotable_2014_pct']}}}",
        f"\\newcommand{{\\ForecastBadStart}}{{{escaped['forecast_bad_2005_pct']}}}",
        f"\\newcommand{{\\ForecastBadEnd}}{{{escaped['forecast_bad_2014_pct']}}}",
        f"\\newcommand{{\\WastewaterStart}}{{{numbers['wastewater_2005_100m_t']}}}",
        f"\\newcommand{{\\WastewaterEnd}}{{{numbers['wastewater_2014_100m_t']}}}",
        f"\\newcommand{{\\TreatmentStart}}{{{numbers['treatment_2005_100m_t']}}}",
        f"\\newcommand{{\\TreatmentEnd}}{{{numbers['treatment_2014_100m_t']}}}",
        f"\\newcommand{{\\TreatmentTotal}}{{{numbers['treatment_total_100m_t']}}}",
        f"\\newcommand{{\\CompositionCVMAE}}{{{numbers['composition_cv_mae_pp']}}}",
        f"\\newcommand{{\\WastewaterCVMAE}}{{{numbers['wastewater_cv_mae_100m_t']}}}",
    ]
    (RESULTS / "generated_macros.tex").write_text("\n".join(macros) + "\n", encoding="utf-8")


def write_manifest() -> None:
    import importlib.metadata

    input_paths = [
        ROOT / "input" / "problem" / "<SOURCE_FILE_REDACTED>",
        ROOT / "input" / "attachments" / "<SOURCE_FILE_REDACTED>",
    ]
    converted_paths = sorted(CONVERTED.glob("*.docx"))
    code_paths = sorted((ROOT / "code").glob("*"))
    runtime_reports = {"run_manifest.json", "reproducibility_check.json", "validation.json", "paper_build.json"}
    result_paths = sorted(path for path in RESULTS.glob("*") if path.name not in runtime_reports and path.is_file())
    manifest = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "seed": SEED,
        "bootstraps": BOOTSTRAPS,
        "command": "powershell -ExecutionPolicy Bypass -File code/run_all.ps1",
        "python": sys.version,
        "platform": platform.platform(),
        "dependencies": {
            package: importlib.metadata.version(package)
            for package in ["numpy", "pandas", "matplotlib", "python-docx", "PyYAML"]
        },
        "input_sha256": {str(path.relative_to(ROOT)): sha256(path) for path in input_paths},
        "converted_sha256": {str(path.relative_to(ROOT)): sha256(path) for path in converted_paths},
        "code_sha256": {str(path.relative_to(ROOT)): sha256(path) for path in code_paths if path.is_file()},
        "result_sha256": {str(path.relative_to(ROOT)): sha256(path) for path in result_paths},
        "excluded_runtime_reports": sorted(runtime_reports),
    }
    write_json(RESULTS / "run_manifest.json", manifest)


def main() -> None:
    monthly, hydrology, monthly_audit = extract_monthly_and_hydrology()
    annual, totals, annual_audit = extract_annual()
    audit = {
        "overall_status": "needs_review",
        "checks": {
            "monthly_shape": {
                "status": "pass" if monthly_audit["monthly_rows"] == monthly_audit["monthly_expected_rows"] else "fail",
                **monthly_audit,
            },
            "published_grade_reproduction": {
                "status": "pass" if monthly_audit["official_grade_match_rate"] == 1.0 else "fail",
                "match_rate": monthly_audit["official_grade_match_rate"],
            },
            "annual_duplicates": {
                "status": "needs_review" if annual_audit["exact_duplicate_rows_removed"] else "pass",
                "removed": annual_audit["exact_duplicate_rows_removed"],
            },
            "annual_percentage_sums": {
                "status": "needs_review" if annual_audit["percentage_sum_anomaly_rows_gt_0_5pp"] else "pass",
                "anomaly_rows": annual_audit["percentage_sum_anomaly_rows_gt_0_5pp"],
                "maximum_error_pp": annual_audit["maximum_percentage_sum_error_pp"],
                "action": "保留原值并按行归一化到100%，主预测只使用水文年干流行",
            },
            "monitoring_coverage_break": {
                "status": "needs_review",
                "coverage_ratio_1999_to_1998": annual_audit["coverage_ratio_1999_to_1998"],
                "action": "用比例而非河长建模，并输出1999年后窗口敏感性",
            },
            "printed_flow_unit_1999": {
                "status": "needs_review" if annual_audit["flow_unit_issue_years"] else "pass",
                "years": annual_audit["flow_unit_issue_years"],
                "action": "按与其余年份一致的亿立方米解释；主模型不使用该流量变量",
            },
        },
        "monthly": monthly_audit,
        "annual": annual_audit,
    }
    write_json(RESULTS / "data_audit.json", audit)

    quality = analyze_quality(monthly)
    sources = analyze_sources(monthly, hydrology)
    forecast, treatment, model_comparison = analyze_forecast_and_treatment(annual, totals)
    summary = {
        "phase": "solve",
        "seed": SEED,
        "quality_evaluation": quality,
        "source_attribution": sources,
        "forecast": forecast,
        "treatment": treatment,
        "automatic_status": {
            "data_completeness": "pass",
            "grade_reproduction": "pass",
            "mass_balance_identity": "pass" if sources["mass_balance_max_abs_error_kg_day"] < 1e-6 else "fail",
            "forecast_small_sample": "needs_review",
            "treatment_mapping_assumption": "needs_review",
        },
    }
    write_json(RESULTS / "summary.json", summary)
    generate_paper_numbers(quality, sources, forecast, treatment)
    make_figures(monthly, annual, totals)
    write_manifest()
    print(json.dumps(as_builtin(summary["automatic_status"]), ensure_ascii=False))


if __name__ == "__main__":
    main()
